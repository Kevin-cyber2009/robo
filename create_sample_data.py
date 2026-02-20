"""
create_sample_data.py - Tạo dữ liệu mẫu
==========================================
Chạy script này để tạo:
1. File .docx mẫu với đủ 3 loại câu hỏi
2. Cấu trúc thư mục data/ mẫu
3. File JSON mẫu (để test ngay không cần docx)

Chạy: python create_sample_data.py
"""

import os
import sys
import json
import uuid

BASE_DIR = os.path.dirname(__file__)
DATA_DIR = os.path.join(BASE_DIR, "data")


def create_sample_json():
    """Tạo file .json mẫu trực tiếp (không cần python-docx)."""
    questions = [
        # === TRẮC NGHIỆM (easy) ===
        {
            "id": str(uuid.uuid4()),
            "type": "multiple_choice",
            "difficulty": "easy",
            "question": "Thủ đô của Việt Nam là?",
            "choices": {"A": "Hà Nội", "B": "Hồ Chí Minh", "C": "Đà Nẵng", "D": "Huế"},
            "answer": "A",
            "passage": "",
            "used": False,
        },
        {
            "id": str(uuid.uuid4()),
            "type": "multiple_choice",
            "difficulty": "easy",
            "question": "Số nguyên tố nhỏ nhất lớn hơn 10 là?",
            "choices": {"A": "11", "B": "12", "C": "13", "D": "15"},
            "answer": "A",
            "passage": "",
            "used": False,
        },
        {
            "id": str(uuid.uuid4()),
            "type": "multiple_choice",
            "difficulty": "easy",
            "question": "Đơn vị đo nhiệt độ trong hệ SI là?",
            "choices": {"A": "Celsius", "B": "Fahrenheit", "C": "Kelvin", "D": "Rankine"},
            "answer": "C",
            "passage": "",
            "used": False,
        },
        {
            "id": str(uuid.uuid4()),
            "type": "multiple_choice",
            "difficulty": "easy",
            "question": "Công thức hóa học của nước là?",
            "choices": {"A": "CO₂", "B": "H₂O", "C": "NaCl", "D": "O₂"},
            "answer": "B",
            "passage": "",
            "used": False,
        },
        # === TRẮC NGHIỆM (medium) ===
        {
            "id": str(uuid.uuid4()),
            "type": "multiple_choice",
            "difficulty": "medium",
            "question": "Định luật II Newton phát biểu rằng?",
            "choices": {
                "A": "Mọi vật đứng yên hoặc chuyển động thẳng đều",
                "B": "F = ma, lực bằng khối lượng nhân gia tốc",
                "C": "Mọi tác dụng đều có phản tác dụng bằng nhau",
                "D": "Năng lượng không tự sinh ra hay mất đi",
            },
            "answer": "B",
            "passage": "",
            "used": False,
        },
        {
            "id": str(uuid.uuid4()),
            "type": "multiple_choice",
            "difficulty": "medium",
            "question": "Quá trình quang hợp xảy ra ở bào quan nào trong tế bào thực vật?",
            "choices": {
                "A": "Ty thể",
                "B": "Ribosome",
                "C": "Lục lạp",
                "D": "Nhân tế bào",
            },
            "answer": "C",
            "passage": "",
            "used": False,
        },
        {
            "id": str(uuid.uuid4()),
            "type": "multiple_choice",
            "difficulty": "medium",
            "question": "Trong lập trình, từ khóa nào dùng để khai báo hàm trong Python?",
            "choices": {"A": "function", "B": "func", "C": "def", "D": "fn"},
            "answer": "C",
            "passage": "",
            "used": False,
        },
        # === TRẢ LỜI NGẮN (easy) ===
        {
            "id": str(uuid.uuid4()),
            "type": "short_answer",
            "difficulty": "easy",
            "question": "7 × 8 = ?",
            "choices": {},
            "answer": "56",
            "passage": "",
            "used": False,
        },
        {
            "id": str(uuid.uuid4()),
            "type": "short_answer",
            "difficulty": "easy",
            "question": "Động vật nào được gọi là 'Chúa tể rừng xanh'?",
            "choices": {},
            "answer": "sư tử",
            "passage": "",
            "used": False,
        },
        # === TRẢ LỜI NGẮN (medium) ===
        {
            "id": str(uuid.uuid4()),
            "type": "short_answer",
            "difficulty": "medium",
            "question": "Căn bậc hai của 144 là?",
            "choices": {},
            "answer": "12",
            "passage": "",
            "used": False,
        },
        {
            "id": str(uuid.uuid4()),
            "type": "short_answer",
            "difficulty": "medium",
            "question": "Chu vi hình tròn có bán kính r được tính bởi công thức nào? (viết công thức)",
            "choices": {},
            "answer": "2πr",
            "passage": "",
            "used": False,
        },
        # === PHÂN TÍCH DỮ KIỆN (hard) ===
        {
            "id": str(uuid.uuid4()),
            "type": "fact_analysis",
            "difficulty": "hard",
            "question": "Xác định đúng/sai cho các nhận định dưới đây:",
            "choices": {
                "A": "Nhiệt độ sôi của nước ở áp suất chuẩn là 100°C",
                "B": "CO₂ là khí giúp thực vật quang hợp",
                "C": "Mặt trăng tự phát ra ánh sáng",
                "D": "Ozone bảo vệ Trái đất khỏi tia UV",
            },
            "answer": {"A": True, "B": True, "C": False, "D": True},
            "passage": "Khoa học tự nhiên nghiên cứu các hiện tượng vật lý và hóa học trong tự nhiên. Các quy luật tự nhiên được kiểm chứng qua thực nghiệm và quan sát.",
            "passage": "Khoa học tự nhiên nghiên cứu các hiện tượng vật lý và hóa học. Nhiệt độ sôi của nước tinh khiết là 100°C tại áp suất 1 atm. Cây xanh sử dụng CO₂ và ánh sáng để tổng hợp glucose. Mặt trăng không tự phát sáng mà phản chiếu ánh sáng mặt trời. Tầng ozone hấp thụ phần lớn tia cực tím (UV) có hại từ mặt trời.",
            "used": False,
        },
        {
            "id": str(uuid.uuid4()),
            "type": "fact_analysis",
            "difficulty": "hard",
            "question": "Xác định đúng/sai cho các nhận định dưới đây:",
            "choices": {
                "A": "Python là ngôn ngữ lập trình thông dịch (interpreted)",
                "B": "Java cần biên dịch ra mã máy trực tiếp như C++",
                "C": "HTML là ngôn ngữ lập trình",
                "D": "Git là hệ thống quản lý phiên bản",
            },
            "answer": {"A": True, "B": False, "C": False, "D": True},
            "passage": "Trong phát triển phần mềm, có nhiều ngôn ngữ lập trình và công cụ khác nhau. Python là ngôn ngữ thông dịch, phổ biến trong khoa học dữ liệu và AI. Java biên dịch ra bytecode chạy trên JVM, không phải mã máy trực tiếp. HTML là ngôn ngữ đánh dấu (markup language), không phải ngôn ngữ lập trình. Git là hệ thống quản lý phiên bản phân tán rất phổ biến.",
            "used": False,
        },
        # === Thêm câu hard ===
        {
            "id": str(uuid.uuid4()),
            "type": "multiple_choice",
            "difficulty": "hard",
            "question": "Trong thuật toán Quick Sort, trường hợp xấu nhất có độ phức tạp thời gian là?",
            "choices": {
                "A": "O(n log n)",
                "B": "O(n²)",
                "C": "O(n)",
                "D": "O(log n)",
            },
            "answer": "B",
            "passage": "",
            "used": False,
        },
        {
            "id": str(uuid.uuid4()),
            "type": "multiple_choice",
            "difficulty": "hard",
            "question": "Nguyên lý nào khẳng định không thể xác định chính xác đồng thời vị trí và vận tốc của một hạt?",
            "choices": {
                "A": "Nguyên lý bảo toàn năng lượng",
                "B": "Nguyên lý bất định Heisenberg",
                "C": "Nguyên lý tương đối Einstein",
                "D": "Nguyên lý Pauli",
            },
            "answer": "B",
            "passage": "",
            "used": False,
        },
    ]

    return questions


def create_sample_docx():
    """Tạo file .docx mẫu với đúng định dạng."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("⚠ Cần cài python-docx để tạo file .docx mẫu.")
        print("  Chạy: pip install python-docx")
        return None

    doc = Document()

    # Tiêu đề
    title = doc.add_heading("BỘ ĐỀ MẪU - RoboLearn Shooter", level=1)

    doc.add_paragraph("File này là ví dụ định dạng câu hỏi chuẩn cho hệ thống.")
    doc.add_paragraph("")

    # === Câu 1: Trắc nghiệm easy ===
    doc.add_paragraph("[MC] [easy]")
    doc.add_paragraph("Câu hỏi: Thủ đô của Việt Nam là?")
    doc.add_paragraph("A. Hà Nội")
    doc.add_paragraph("B. Hồ Chí Minh")
    doc.add_paragraph("C. Đà Nẵng")
    doc.add_paragraph("D. Huế")
    doc.add_paragraph("Đáp án: A")
    doc.add_paragraph("")

    # === Câu 2: Trắc nghiệm medium ===
    doc.add_paragraph("[MC] [medium]")
    doc.add_paragraph("Câu hỏi: Quá trình quang hợp xảy ra ở bào quan nào?")
    doc.add_paragraph("A. Ty thể")
    doc.add_paragraph("B. Ribosome")
    doc.add_paragraph("C. Lục lạp")
    doc.add_paragraph("D. Nhân tế bào")
    doc.add_paragraph("Đáp án: C")
    doc.add_paragraph("")

    # === Câu 3: Trả lời ngắn ===
    doc.add_paragraph("[SA] [easy]")
    doc.add_paragraph("Câu hỏi: 7 × 8 = ?")
    doc.add_paragraph("Đáp án: 56")
    doc.add_paragraph("")

    # === Câu 4: Trả lời ngắn medium ===
    doc.add_paragraph("[SA] [medium]")
    doc.add_paragraph("Câu hỏi: Căn bậc hai của 144 là?")
    doc.add_paragraph("Đáp án: 12")
    doc.add_paragraph("")

    # === Câu 5: Phân tích dữ kiện ===
    doc.add_paragraph("[FA] [hard]")
    doc.add_paragraph("Dữ kiện: Python là ngôn ngữ thông dịch phổ biến. Java biên dịch ra bytecode chạy trên JVM. HTML là ngôn ngữ đánh dấu. Git là hệ thống quản lý phiên bản phân tán.")
    doc.add_paragraph("A. Python là ngôn ngữ lập trình thông dịch")
    doc.add_paragraph("B. Java biên dịch ra mã máy trực tiếp như C++")
    doc.add_paragraph("C. HTML là ngôn ngữ lập trình")
    doc.add_paragraph("D. Git là hệ thống quản lý phiên bản")
    doc.add_paragraph("Đáp án: A-Đúng,B-Sai,C-Sai,D-Đúng")
    doc.add_paragraph("")

    # === Câu 6: Trắc nghiệm hard ===
    doc.add_paragraph("[MC] [hard]")
    doc.add_paragraph("Câu hỏi: Nguyên lý bất định Heisenberg phát biểu rằng không thể xác định chính xác đồng thời điều gì?")
    doc.add_paragraph("A. Khối lượng và điện tích")
    doc.add_paragraph("B. Vị trí và động lượng")
    doc.add_paragraph("C. Nhiệt độ và áp suất")
    doc.add_paragraph("D. Tốc độ và gia tốc")
    doc.add_paragraph("Đáp án: B")

    return doc


def main():
    print("=" * 60)
    print("  RoboLearn Shooter - Tạo Dữ Liệu Mẫu")
    print("=" * 60)

    # Tạo cấu trúc thư mục mẫu
    sample_class = "Lớp 10"
    sample_subject = "Tổng Hợp"
    subject_dir = os.path.join(DATA_DIR, sample_class, sample_subject)
    os.makedirs(subject_dir, exist_ok=True)
    print(f"✓ Tạo thư mục: data/{sample_class}/{sample_subject}/")

    # Lưu file JSON mẫu
    questions = create_sample_json()
    json_path = os.path.join(subject_dir, "bo_de_mau.json")
    with open(json_path, "w", encoding="utf-8") as f:
        json.dump({
            "metadata": {
                "source_file": "bo_de_mau.docx",
                "class": sample_class,
                "subject": sample_subject,
            },
            "questions": questions,
        }, f, ensure_ascii=False, indent=2)
    print(f"✓ Tạo file JSON mẫu: {json_path}")
    print(f"  → {len(questions)} câu hỏi ({sum(1 for q in questions if q['difficulty']=='easy')} dễ, "
          f"{sum(1 for q in questions if q['difficulty']=='medium')} trung bình, "
          f"{sum(1 for q in questions if q['difficulty']=='hard')} khó)")

    # Tạo file .docx mẫu (nếu python-docx có sẵn)
    docx_dir = os.path.join(BASE_DIR, "examples")
    os.makedirs(docx_dir, exist_ok=True)
    doc = create_sample_docx()
    if doc:
        docx_path = os.path.join(docx_dir, "bo_de_mau.docx")
        doc.save(docx_path)
        print(f"✓ Tạo file .docx mẫu: {docx_path}")
    else:
        print("⚠ Bỏ qua tạo file .docx (không có python-docx)")

    print()
    print("✅ Sẵn sàng! Chạy game bằng: python main.py")
    print("   Dữ liệu mẫu đã được tạo tại: data/Lớp 10/Tổng Hợp/")


if __name__ == "__main__":
    main()
